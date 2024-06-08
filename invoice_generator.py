import logging
import pandas as pd
from docx import Document
from datetime import datetime
import calendar
import json
import os


with open('config.json', 'r') as config_file:
    config = json.load(config_file)

CSV_PATH = config['csv_path']
TEMPLATE_FILE = config['template_file']
OUTPUT_DIR = config['output_dir']
DAY_RATE_PER_WEEK = config['day_rate_per_week']
CLIENT = config['client']
ACCOUNT_HOLDER = config['account_holder']
COMPANY_NAME = config['company_name']
COMPANY_ADDRESS = config['company_address']
COMPANY_EMAIL = config['company_email']

def load_dataframe(csv_path):
    """Load the CSV file into a DataFrame."""
    try:
        return pd.read_csv(csv_path, sep='\t')
    except (FileNotFoundError, pd.errors.ParserError) as e:
        logging.error(f"Error loading DataFrame from {csv_path}: {str(e)}")
        return None

def format_date(date):
    """Format date to 'Month Day'."""
    return date.strftime("%B %-d")

def get_weekday_ranges_for_csv_month(df, year=None):
    """
    Get weekday ranges for the month extracted from the CSV DataFrame.

    Args:
    - df: DataFrame containing CSV data.
    - year: Year to use for month calculation. Defaults to current year if None.

    Returns:
    - weekday_ranges: List of tuples representing start and end dates of each week.
    - formatted_ranges: List of formatted strings for each weekday range.
    """
    month_str = df.iloc[0, 1]  # Assuming month is in the second column of the first row
    month = datetime.strptime(month_str, '%B').month
    if year is None:
        year = datetime.now().year

    first_day_of_month = datetime(year, month, 1)
    last_day_of_month = datetime(year, month, calendar.monthrange(year, month)[1])
    
    weekday_ranges = []
    weeks = calendar.monthcalendar(year, month)
    
    for week in weeks:
        week_days = [day for day in week if day != 0]
        if not week_days:
            continue

        week_start = datetime(year, month, week_days[0])
        week_end = datetime(year, month, week_days[-1])
        if week_start < first_day_of_month:
            week_start = first_day_of_month
        if week_end > last_day_of_month:
            week_end = last_day_of_month
        weekday_ranges.append((week_start, week_end))

    formatted_ranges = [f"{format_date(start)} - {format_date(end)}" for (start, end) in weekday_ranges]
    
    return weekday_ranges, formatted_ranges

def count_workdays_in_week(df, start_date, end_date):
    count = 0
    for day in range(start_date.day, end_date.day + 1):
        if df[str(day)].values[0] == 8:
            count += 1
    return count

def create_data_dictionary(invoice_number, formatted_date, ACCOUNT_HOLDER, COMPANY_NAME, COMPANY_ADDRESS, COMPANY_EMAIL):
    return {
        '{{INVOICE_NUMBER}}': invoice_number,
        '{{DATE}}': formatted_date,
        '{{ACCOUNT_HOLDER}}': ACCOUNT_HOLDER,
        '{{ROUTING_NUMBER}}': config['routing_number'],
        '{{SWIFT_BIC}}': config['swift_bic'],
        '{{ACCOUNT_NUMBER}}': config['account_number'],
        '{{WISE_ADDRESS}}': config['wise_address'],
        '{{COMPANY_NAME}}': COMPANY_NAME,
        '{{COMPANY_ADDRESS}}': COMPANY_ADDRESS,
        '{{COMPANY_EMAIL}}': COMPANY_EMAIL
    }

def fill_template(template_file, output_path, data):
    if not os.path.exists(template_file):
        logging.error(f"The file {template_file} was not found.")
        return
    
    doc = Document(template_file)
    for key, value in data.items():
        for paragraph in doc.paragraphs:
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(value))
    doc.save(output_path)
    
def generate_invoice(df, invoice_number, formatted_date):
    weekday_ranges, formatted_ranges = get_weekday_ranges_for_csv_month(df)
    for index, row in df.iterrows():
        data = create_data_dictionary(invoice_number, formatted_date, ACCOUNT_HOLDER, COMPANY_NAME, COMPANY_ADDRESS, COMPANY_EMAIL)
        total_due = 0
        workdays_cache = {}
        
        for i, period in enumerate(weekday_ranges):
            start, end = period
            if period not in workdays_cache:
                workdays_cache[period] = count_workdays_in_week(df, start, end)
            workdays = workdays_cache[period]
            week_description = f"{workdays} Days - {CLIENT}"
            data[f'{{{{WORK_PERIOD{i+1}}}}}'] = formatted_ranges[i]
            data[f'{{{{DESCRIPTION{i+1}}}}}'] = week_description
            data[f'{{{{DAY_RATE{i+1}}}}}'] = f'${DAY_RATE_PER_WEEK:.2f}'
            total = workdays * DAY_RATE_PER_WEEK
            data[f'{{{{TOTAL{i+1}}}}}'] = f'${total:.2f}'
            total_due += total
        
        data['{{TOTAL_DUE}}'] = f'${total_due:.2f}'
        
        output_path = os.path.join(OUTPUT_DIR, f'invoice_{ACCOUNT_HOLDER}_{invoice_number}.docx')
        fill_template(TEMPLATE_FILE, output_path, data)

def main():
    try:
        with open('config.json', 'r') as config_file:
            config = json.load(config_file)
    except json.JSONDecodeError as e:
        logging.error(f"Error parsing configuration JSON: {str(e)}")
        return
    
    df = load_dataframe(CSV_PATH)
    if df is not None:
        current_date = datetime.now()
        invoice_number = current_date.strftime("%Y%m")
        formatted_date = current_date.strftime("%d %b %Y")
        generate_invoice(df, invoice_number, formatted_date)
    else:
        logging.error(f"The file {config['template_file']} was not found.")

if __name__ == "__main__":
    main()